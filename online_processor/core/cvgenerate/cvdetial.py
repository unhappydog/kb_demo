from docx.shared import Pt
import re
from core.cvgenerate.util import geUtil

class detailinfo(object):
    def base(self,document,data):
        table = document.add_table(rows=2, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '期望从事职业'
        hdr_cells[1].text = data['expectedOccupation']
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = '简历更新时间'
        hdr_cells[1].text = data['updateTime'].strftime("%Y-%m-%d")
        cvid = document.add_paragraph("ID：")
        cvid.add_run(data['_id']).bold = True
        paragraph = document.add_paragraph('')
        run = paragraph.add_run(data['name'])
        run.bold = True
        run.font.size = Pt(16)
        cvbase = document.add_paragraph('')
        cvbase.add_run(data['gender'] + "    ").bold = True
        cvbase.add_run(str(data['age']) + "岁" + "    ").bold = True
        date = "".join(re.findall('\d{4}-\d{2}', data['birthday'].strftime("%Y-%m"))).replace('-', '年')
        cvbase.add_run(date + "月" + "    ").bold = True
        cvbase.add_run(str(data['workYear']) + "工作经验" + "    ").bold = True
        cvbase.add_run(data['highestEducationBackground'] + "  " + data['marital']).bold = True
        cvbase1 = document.add_paragraph('')
        cvbase1.add_run('现居住地：' + data['currentAddress'])
        cvbase1.add_run("|" + "户口：" + data['domicilePlace'] + "|")
        cvbase1.add_run(data['politicsStatus'])
        print('已经运行')

    def intent(self, document, data,title):
        intentdata = {'expectedWorkplace': '期望工作地区', 'expectedSalary': '期望薪资',
                      'expectedStatus': '目前工作状况', 'expectedWorkNature': '期望工作性质', 'expectedOccupation': '期望从事职业',
                      'expectedIndustry': '期望从事行业'}
        geUtil().title_style(title,document)
        indata = {}
        rows = 6
        table = document.add_table(rows=rows, cols=2)
        table.autofit = True
        table.cell(0, 1).width = Pt(700)
        for i in data:
            for k, v in intentdata.items():
                if i == k:
                    indata[v] = data[i]
        for i, k in zip(range(rows), indata):
            indent_cells = table.rows[i].cells
            indent_cells[0].text = k + "："
            indent_cells[1].text = indata[k]
    def selfEvaluation(self,document, data,title):
        if data['selfEvaluation']:
            geUtil().title_style(title,document)
            if '\n' in data['selfEvaluation']:
                sellfsplit = re.findall('.*?\n', data['selfEvaluation'])
                selfstr = ''
                for se in sellfsplit:
                    if len(se) > 1:
                        selfstr += se
                document.add_paragraph(selfstr)
            else:
                document.add_paragraph(data['selfEvaluation'])

    def workExperience(self, document, data,title):
        workdata = data['workExperience']
        if workdata:
            geUtil().title_style(title,document)
            for i in workdata:
                toptable, paragraph, starttime, endtime = geUtil().time(document, i, data['updateTime'])
                paragraph.add_run(starttime + '-' + endtime + '\t' + i['workCompany'] + '\t' + "（" + i[
                    'workTimePeriod'] + "）").bold = True
                if i['workPosition']:
                    toptable.add_row()
                    nexttabcell = toptable.cell(1, 0)
                    para2 = nexttabcell.paragraphs[0]
                    para2run = para2.add_run(i['workPosition'])
                    para2run.bold = True
                    try:
                        if i['workSalary']!=' ' or i['workSalary']!='':
                                para2.add_run("|" + i['workSalary']).bold = True
                                para2run.bold = True
                                para2run.font.size = Pt(10)
                    except:
                        para2.add_run('')
                try:
                    if i['workCompanyIndustry']:
                            toptable.add_row()
                            nexttabcell = toptable.cell(2, 0)
                            para2 = nexttabcell.paragraphs[0]
                            para2.add_run(i['workCompanyIndustry'])
                            try:
                                    if i['workCompanyNature']:
                                        para2.add_run("|"+"企业性质："+i['workCompanyIndustry'])

                                    if i['workCompanyScale']:
                                        para2.add_run("|"+"规模："+i['workCompanyScale'])
                            except:
                                para2.add_run("")
                except:
                    print('')
                if i['workDescription']:
                    subtable = document.add_table(rows=1, cols=2)
                    subtable.autofit = True
                    subtable.cell(0, 1).width = Pt(700)
                    subtable.cell(0, 0).paragraphs[0].add_run('工作描述：')
                    if '\n' in i['workDescription']:
                        wordtext=re.sub('[ \t]','',i['workDescription'])
                    else:
                        wordtext =re.sub('[ \t]','', re.sub('[；。]', '；\n', i['workDescription']))
                    subtable.cell(0, 1).paragraphs[0].add_run(wordtext)

    def projectExperience(self, document, data,title):
        projectdata = data['projectExperience']
        if projectdata:
            geUtil().title_style(title,document)
            prodict = {'projectSoftwareEnv': '软件环境', 'projectHardwareEnv': '硬件环境', 'projectTool': '开发工具',
                       'projectDuty': '责任描述', 'projectDescription': '项目描述'}
            for pro in projectdata:
                toptable,paragraph, starttime, endtime = geUtil().time(document, pro, data['updateTime'])
                try:
                    if pro['projectTimePeriod']:
                        paragraph.add_run(
                            starttime + '-' + endtime + '\t' + pro['projectName']+'\t'+'（'+pro['projectTimePeriod']+'）').bold = True
                except:
                    paragraph.add_run(
                        starttime + '-' + endtime + '\t' + pro['projectName']).bold = True
                prodata = {}
                count=0
                for i in pro:
                    for k, v in prodict.items():
                        if i == k:
                            if pro[i]!='':
                                 prodata[v] = pro[i]
                            else:
                                count+=1
                try:
                    record=len(prodict) -count-1
                    subtable = document.add_table(rows=record, cols=2)
                    subtable.autofit = True
                    subtable.cell(0, 1).width = Pt(700)
                    for i, k in zip(range(record), prodata):
                        indent_cells = subtable.rows[i].cells
                        indent_cells[0].text = k+"："
                        indent_cells[1].text = re.sub('[\t]','',prodata[k])
                except:
                    record = len(prodict) - count
                    subtable = document.add_table(rows=record, cols=2)
                    subtable.autofit = True
                    subtable.cell(0, 1).width = Pt(700)
                    for i, k in zip(range(record), prodata):
                        indent_cells = subtable.rows[i].cells
                        indent_cells[0].text = k + "："
                        indent_cells[1].text = re.sub('[\t]', '', prodata[k])

    def educationExperience(self, document, data,title):
        eductiondata = data['educationExperience']
        geUtil().title_style(title,document)
        for edu in eductiondata:
            toptable,paragraph, starttime, endtime =geUtil().time(document, edu,data['updateTime'])
            paragraph.add_run(
                starttime + '-' + endtime + '\t\t' + edu['educationSchool'] + '\t' + edu['educationMajor'] + '\t' + edu[
                    'educationDegree'])

    def language(self, document, data,title):
        language = data['language']
        if language:
            geUtil().title_style(title,document)
            for lau in language:
                lautext = document.add_paragraph('')
                lautext.add_run(lau['languageKind'] + '\t' + '读写程度' + lau['readingWritingLevel'] + "|" + '听说程度' + lau[
                    'listeningSpeakingLevel'])

    def skill(self, document, data,title):
        skill = data['skill']
        if skill:
            geUtil().title_style(title, document)
            toptable = document.add_table(rows=0, cols=1)
            if isinstance(skill,list):
                for sk in skill:
                    toptable.add_row()
                    toptablecells = toptable.cell(skill.index(sk), 0)
                    paragraph = toptablecells.paragraphs[0]
                    paragraph.add_run(sk['name'] + "：" + sk['skillMastery'])
            else:
                toptablecells = toptable.cell(skill.index(0), 0)
                paragraph = toptablecells.paragraphs[0]
                paragraph.add_run(skill)
    def trainingExperience(self,document,data,title):
        traindata=data['trainingExperience']
        tradict={'trainingOrg':'培训机构','traininigPlace':'培训地点','trainingDescription':'培训描述','trainingCertificate':'所获证书'}
        if traindata:
            geUtil().title_style(title,document)
            for tra in traindata:
                toptable, paragraph, starttime, endtime = geUtil().time(document, tra, data['updateTime'])
                paragraph.add_run(
                    starttime + '-' + endtime + '\t\t' + tra['trainingCourse']).bold = True
                prodata = {}
                count = 0
                for i in tra:
                    for k, v in tradict.items():
                        if i == k:
                            if tra[i] != '':
                                prodata[v] = tra[i]
                            else:
                                count += 1

                subtable = document.add_table(rows=len(tradict) - count, cols=2)
                subtable.autofit = True
                subtable.cell(0, 1).width = Pt(700)
                for i, k in zip(range(len(tradict) - count), prodata):
                    indent_cells = subtable.rows[i].cells
                    indent_cells[0].text = k
                    indent_cells[1].text = prodata[k]

    def certificate(self,document,data,title):
        cerdata=data['certificate']
        if cerdata:
            geUtil().title_style(title, document)
            toptable = document.add_table(rows=0, cols=1)
            for cer in cerdata:
                toptable.add_row()
                toptablecells = toptable.cell(cerdata.index(cer), 0)
                paragraph = toptablecells.paragraphs[0]
                time=cer['time'].strftime("%Y.%m")
                paragraph.add_run(time+"\t"+cer['name'])

    def hobby(self, document, data,title):
        hobby = data['hobby']
        if hobby:
            geUtil().title_style(title,document)
            document.add_paragraph(hobby)