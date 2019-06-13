from data_access.controller.CVController4Mongo import CVController4Mongo
from docx.shared import RGBColor,Pt
# from core.cvgenerate.detailcv import detailcvinfo
class AMetaclass(type):

    def __new__(cls, name, bases, attrs, *args, **kwargs):
        count = 0
        attrs["__Func__"] = []
        for k, v in attrs.items():
            if "get_" in k:
                attrs["__Func__"].append(k)
                count += 1
        attrs["__FuncCount__"] = count
        return type.__new__(cls, name, bases, attrs)


class geUtil(metaclass=AMetaclass):
    def getdata(self):
        data = CVController4Mongo().get_datas()
        # data = CVController4Mongo().get_data_by_id(_id='(ZXuD0v0F4JT3yw(GUfC3Q')
        return data
    def title_style(self,title,document):
        titlestyle = document.add_paragraph('', style='Title')
        titlerun = titlestyle.add_run(title)
        titlerun.font.color.rgb = RGBColor(54, 95, 145)
        titlerun.font.size=Pt(16)
        titlerun.bold = True

    def function_name(self,fun):
        funlist = []
        function = fun.__dir__()
        for fc in function:
            if '__' not in fc and fc != 'firstflood' and fc != 'function_name':
                funlist.append(fc)
        return funlist
    def time(self,document,vari,data):
        start=''
        end=''
        for da in vari:
            if 'StartTime' in da:
                start=da
            if 'EndTime' in da:
                end=da
        toptable = document.add_table(rows=1, cols=1)
        toptablecells = toptable.cell(0, 0)
        paragraph = toptablecells.paragraphs[0]
        starttime = vari[start].strftime("%Y.%m")
        if vari[end]:
            endtime = vari[end].strftime("%Y.%m")
        else:
            endtime = data.strftime("%Y.%m")
        return toptable,paragraph,starttime,endtime

    def en_to_cn(self,temp):
        return {
            'intent': '求职意向',
            'selfEvaluation': '自我评价',
            'workExperience': '工作经历',
            'projectExperience': '项目经历',
            'educationExperience': '教育经历',
            'certificate': '证书',
            'trainingExperience':'培训经历',
            'award': '获奖状况',
            'associationExperience': '在校实践情况',
            'language': '语言能力',
            'skill': '专业技能',
            'hobby': '兴趣爱好'
        }.get(temp, temp)


if __name__=='__main__':
    print(geUtil().getdata())