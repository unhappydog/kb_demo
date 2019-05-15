import re
import docx
from core.cvprase.Util import adjunction as ad
import subprocess
import os
class util(object):
    def doc2docx(self,root):
        output = subprocess.check_output(
            ["soffice", "--headless", "--convert-to", "docx:MS Word 2007 XML", root,
             "--outdir",os.path.split(root)[0]])

    """字典整合"""
    def integration(self, dict1, dict2):
        return {**dict1, **dict2}

    def listsplit(self,datalist):
        listpro = []
        for tra in datalist:
            if "".join(re.compile('\d{4}.\d{2} - 至今').findall("".join(tra))):
                listpro.append(datalist.index(tra))
            elif "".join(re.compile('\d{4}.\d{2} - \d{4}.\d{2}').findall("".join(tra))):
                listpro.append(datalist.index(tra))
        return listpro

    def get_doc(self,path):
        doc = docx.Document(path)
        document = ad.opendocx(path)
        fullText = ad.getdocumenttext(document)
        index=[]
        for i in fullText:
            if '男 ' in i or '女 ' in i:
                index.append(fullText.index(i))
        minindex=min(index)
        name =fullText[minindex-1].strip(' ')
        table = doc.tables
        para = doc.paragraphs
        return table, para, name,fullText

    def tabletxt(self, table):
        c = []
        for row in table.rows:
            for cell in row.cells:
                if cell.text not in c:
                    c.append(cell.text)
        return c

    def tabletxtb(self, table):
        c = []
        try:
            for row in table.rows:
                z = []
                for cell in row.cells:
                    if cell.text not in z:
                        z.append(cell.text)
                c.append(z)
        except:
            for j in table:
                for row in j.rows:
                    z = []
                    for cell in row.cells:
                        if cell.text not in z:
                            z.append(cell.text)
                    c.append(z)

        return c